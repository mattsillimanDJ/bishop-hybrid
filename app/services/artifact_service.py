import csv
import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font


DEFAULT_ARTIFACT_DIR = "/tmp/bishop_artifacts"

DOCX_PHRASES = (
    "make this a word doc",
    "make this a word document",
    "export this as docx",
    "make this a docx",
    "turn this into a word document",
)

XLSX_PHRASES = (
    "make this an excel file",
    "make this a excel file",
    "export this as xlsx",
    "make this a spreadsheet",
    "make this an excel spreadsheet",
    "turn this into a spreadsheet",
)


@dataclass(frozen=True)
class ArtifactExportRequest:
    kind: str
    content: str


@dataclass(frozen=True)
class ArtifactResult:
    kind: str
    path: Path
    filename: str


def get_artifact_dir() -> Path:
    return Path(os.getenv("BISHOP_ARTIFACT_DIR", DEFAULT_ARTIFACT_DIR)).expanduser()


def normalize_artifact_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip()).casefold()


def extract_content_after_phrase(message: str, phrase: str) -> str:
    match = re.search(re.escape(phrase), message or "", re.IGNORECASE)
    if not match:
        return ""

    content = (message or "")[match.end():]
    content = re.sub(r"^\s*[:,-]\s*", "", content)
    return content.strip()


def detect_artifact_export_request(message: str) -> ArtifactExportRequest | None:
    normalized = normalize_artifact_text(message)
    if not normalized:
        return None

    for phrase in DOCX_PHRASES:
        if phrase in normalized:
            return ArtifactExportRequest(
                kind="docx",
                content=extract_content_after_phrase(message, phrase),
            )

    for phrase in XLSX_PHRASES:
        if phrase in normalized:
            return ArtifactExportRequest(
                kind="xlsx",
                content=extract_content_after_phrase(message, phrase),
            )

    return None


def safe_filename(kind: str, now: datetime | None = None) -> str:
    timestamp = (now or datetime.now(timezone.utc)).strftime("%Y%m%d_%H%M%S")
    extension = "docx" if kind == "docx" else "xlsx"
    return f"bishop_artifact_{timestamp}.{extension}"


def title_from_content(content: str) -> str:
    for line in (content or "").splitlines():
        stripped = line.strip().strip("#").strip()
        if stripped:
            return stripped[:80]
    return "Bishop Export"


def create_docx_artifact(content: str, output_dir: Path | None = None) -> ArtifactResult:
    if not (content or "").strip():
        raise ValueError("DOCX artifact content is required.")

    directory = output_dir or get_artifact_dir()
    directory.mkdir(parents=True, exist_ok=True)
    filename = safe_filename("docx")
    path = directory / filename

    document = Document()
    title = title_from_content(content)
    document.add_heading(title, level=0)

    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line == title:
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            document.add_heading(heading_match.group(2).strip(), level=level)
            continue

        if line.endswith(":") and len(line) <= 90:
            document.add_heading(line.rstrip(":"), level=1)
            continue

        bullet_match = re.match(r"^[-*•]\s+(.+)$", line)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue

        numbered_match = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered_match:
            document.add_paragraph(numbered_match.group(1).strip(), style="List Number")
            continue

        document.add_paragraph(line)

    document.save(path)
    return ArtifactResult(kind="docx", path=path, filename=filename)


def is_markdown_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells if cell.strip())


def parse_pipe_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if "|" not in line:
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if is_markdown_separator_row(cells):
            continue
        if any(cells):
            rows.append(cells)
    return rows


def parse_delimited_rows(content: str) -> list[list[str]]:
    lines = [line.strip() for line in (content or "").splitlines() if line.strip()]
    if not lines:
        return []

    pipe_rows = parse_pipe_table(lines)
    if pipe_rows:
        return pipe_rows

    delimiter = "\t" if any("\t" in line for line in lines) else ","
    if any(delimiter in line for line in lines):
        return [
            [cell.strip() for cell in row]
            for row in csv.reader(lines, delimiter=delimiter)
            if any(cell.strip() for cell in row)
        ]

    rows = [["Item"]]
    for line in lines:
        cleaned = re.sub(r"^\s*(?:[-*•]|\d+[.)])\s*", "", line).strip()
        if cleaned:
            rows.append([cleaned])
    return rows


def create_xlsx_artifact(content: str, output_dir: Path | None = None) -> ArtifactResult:
    if not (content or "").strip():
        raise ValueError("XLSX artifact content is required.")

    rows = parse_delimited_rows(content)
    if not rows:
        raise ValueError("XLSX artifact rows are required.")

    directory = output_dir or get_artifact_dir()
    directory.mkdir(parents=True, exist_ok=True)
    filename = safe_filename("xlsx")
    path = directory / filename

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Bishop Export"

    for row in rows:
        worksheet.append(row)

    for cell in worksheet[1]:
        cell.font = Font(bold=True)

    for column_cells in worksheet.columns:
        max_length = max(len(str(cell.value or "")) for cell in column_cells)
        worksheet.column_dimensions[column_cells[0].column_letter].width = min(
            max(max_length + 2, 12),
            42,
        )

    workbook.save(path)
    return ArtifactResult(kind="xlsx", path=path, filename=filename)


def create_artifact(kind: str, content: str, output_dir: Path | None = None) -> ArtifactResult:
    if kind == "docx":
        return create_docx_artifact(content=content, output_dir=output_dir)
    if kind == "xlsx":
        return create_xlsx_artifact(content=content, output_dir=output_dir)
    raise ValueError(f"Unsupported artifact kind: {kind}")
